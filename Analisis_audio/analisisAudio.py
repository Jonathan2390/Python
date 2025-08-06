import os
import math
import time
import threading
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from pydub import AudioSegment
from pydub.playback import play
import whisper
from docx import Document
from PIL import Image


# === Configuración inicial ===
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

app = ctk.CTk()
# app.iconbitmap("icons/ico.ico")
app.title("Análisis de audio")
app.geometry("1100x600")
app.resizable(True, True)

# === Cargar modelo Whisper ===
modelo = whisper.load_model("small", device="cpu")

# === Variables globales ===
archivo_audio = None
frases_clave = []
fragmentos_audio = []
cronometro_activo = False
inicio_tiempo = 0
modo_oscuro = False

# === Cargar íconos ===
'''iconos = {
    "audio": ctk.CTkImage(Image.open("icons/audio.png"), size=(24, 24)),
    "key": ctk.CTkImage(Image.open("icons/key.png"), size=(24, 24)),
    "search": ctk.CTkImage(Image.open("icons/search.png"), size=(24, 24)),
    "export": ctk.CTkImage(Image.open("icons/export.png"), size=(24, 24)),
    "light": ctk.CTkImage(Image.open("icons/light.png"), size=(24, 24)),
    "dark": ctk.CTkImage(Image.open("icons/dark.png"), size=(24, 24)),
    "clean": ctk.CTkImage(Image.open("icons/clean.png"), size=(24, 24)),
    "stop": ctk.CTkImage(Image.open("icons/stop.png"), size=(24, 24))
}'''

# === Elementos de la GUI ===
estado_label = ctk.CTkLabel(
    app, text="🎧 Esperando archivo de audio", text_color="blue",
    font=ctk.CTkFont(size=14), fg_color="transparent", corner_radius=20, width=880)
estado_label.pack(pady=5)

cronometro_label = ctk.CTkLabel(
    app, text="⏱ Tiempo transcurrido: 00:00",
    font=ctk.CTkFont(size=14, weight="bold"))
cronometro_label.pack()

lista_resultados = tk.Listbox(
    app, width=110, height=12, font=("Noto Sans", 11))
lista_resultados.pack(pady=10)
lista_resultados.bind("<Double-Button-1>", lambda e: reproducir_fragmento())

log_texto = ctk.CTkTextbox(app, height=100, width=780)
log_texto.pack(pady=5)
log_texto.configure(
    state="disabled",
    fg_color="transparent",
    corner_radius=20,
    border_width=1
)


def actualizar_cronometro():
    if cronometro_activo:
        tiempo_actual = int(time.time() - inicio_tiempo)
        minutos = tiempo_actual // 60
        segundos = tiempo_actual % 60
        cronometro_label.configure(
            text=f"⏱ Tiempo transcurrido: {minutos:02}:{segundos:02}")
        app.after(1000, actualizar_cronometro)


def log(mensaje):
    log_texto.configure(state="normal")
    log_texto.insert("end", mensaje + "\n")
    log_texto.see("end")
    log_texto.configure(state="disabled")


def seleccionar_audio():
    global archivo_audio
    archivo = filedialog.askopenfilename(
        filetypes=[("Audio .m4a", "*.m4a *.wav *.wma")])
    if archivo:
        archivo_audio = archivo
        estado_label.configure(
            text=f"📁 Archivo cargado: {os.path.basename(archivo)}")
        log(f"✅ Archivo cargado: {archivo}")


def pedir_frases_clave():
    global frases_clave
    entrada = simpledialog.askstring(
        "Frases clave", "Introduce hasta 10 frases clave, separadas por coma:")
    if entrada:
        frases_clave = [f.strip().lower() for f in entrada.split(",")][:10]
        log(f"🔍 Frases clave: {frases_clave}")


def detectar_fragmentos():
    global fragmentos_audio, cronometro_activo, inicio_tiempo

    if not archivo_audio:
        messagebox.showwarning("Aviso", "Selecciona un archivo primero.")
        return
    if not frases_clave:
        messagebox.showwarning("Aviso", "Introduce al menos una frase clave.")
        return

    estado_label.configure(text="🧠 Procesando fragmentos...")
    app.update_idletasks()

    cronometro_label.configure(text="⏱ Tiempo transcurrido: 00:00")
    cronometro_activo = True
    inicio_tiempo = time.time()
    app.after(1000, actualizar_cronometro)

    audio = AudioSegment.from_file(archivo_audio)
    duracion_seg = math.ceil(audio.duration_seconds)
    ventana = 20
    paso = 15

    fragmentos_audio.clear()
    lista_resultados.delete(0, "end")

    for inicio in range(0, duracion_seg, paso):
        if not cronometro_activo:
            log("🛑 Búsqueda interrumpida.")
            break

        fin = min(duracion_seg, inicio + ventana)
        fragmento = audio[inicio * 1000:fin * 1000]

        temp_path = "temp_fragment.wav"
        fragmento.export(temp_path, format="wav")

        try:
            resultado = modelo.transcribe(
                temp_path, language="es", fp16=False, verbose=False)
        except Exception as e:
            log(f"❌ Error al transcribir: {e}")
            continue

        texto = resultado["text"].lower()
        if any(frase in texto for frase in frases_clave):
            ini = max(0, inicio - 1.5) * 1000
            fin_corte = min(duracion_seg, fin + 1.5) * 1000
            recorte = audio[int(ini):int(fin_corte)]
            fragmentos_audio.append(recorte)
            lista_resultados.insert(
                tk.END, f"{time_format(inicio)} - {texto.strip()}")

    # Al finalizar el bucle (ya sea completo o interrumpido):
    cronometro_activo = False

    if fragmentos_audio:
        estado_label.configure(
            text=f"✅ Se encontraron {len(fragmentos_audio)} fragmentos")
        log(f"🔍 Se encontraron {len(fragmentos_audio)} fragmentos relevantes.")
    else:
        estado_label.configure(text="❌ No se encontraron frases.")
        log("⚠️ No se encontraron coincidencias en el audio.")


def reproducir_fragmento():
    seleccion = lista_resultados.curselection()
    if seleccion:
        idx = seleccion[0]
        fragmento = fragmentos_audio[idx]

        def play_thread():
            play(fragmento)

        threading.Thread(target=play_thread, daemon=True).start()


def exportar_a_word():
    if not fragmentos_audio:
        messagebox.showinfo("Sin datos", "No hay fragmentos para exportar.")
        return

    # Preguntar al usuario la ubicación y el nombre del archivo
    nombre_archivo = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Documentos de Word", "*.docx")],
        title="Guardar archivo como",
        initialfile="fragmentos_detectados.docx"
    )

    if not nombre_archivo:
        return  # El usuario canceló el diálogo

    # Crear el documento
    doc = Document()
    doc.add_heading('Fragmentos Detectados', level=1)
    doc.add_paragraph(f"Archivo de origen: {os.path.basename(archivo_audio)}")
    doc.add_paragraph(f"Frases clave usadas: {', '.join(frases_clave)}")
    doc.add_paragraph(" ")

    for i, linea in enumerate(lista_resultados.get(0, tk.END)):
        if " - " in linea:
            tiempo, texto = linea.split(" - ", 1)
            doc.add_heading(f"Fragmento {i + 1} – {tiempo}", level=2)
            doc.add_paragraph(texto)

    # Guardar el archivo
    doc.save(nombre_archivo)
    log(f"📄 Resultados exportados a {nombre_archivo}")
    messagebox.showinfo("Exportación exitosa",
                        f"Archivo guardado como:\n{nombre_archivo}")


def time_format(segundos):
    minutos = int(segundos // 60)
    seg = int(segundos % 60)
    return f"{minutos:02}:{seg:02}"


def limpiar_todo():
    if messagebox.askyesno("Confirmación", "¿Seguro que quieres borrar todo?"):
        lista_resultados.delete(0, "end")
        log_texto.configure(state="normal")
        log_texto.delete("1.0", "end")
        log_texto.configure(state="disabled")
        estado_label.configure(
            text="🎧 Esperando archivo .m4a", fg_color="white")
        fragmentos_audio.clear()
        log("🧹 Todo ha sido limpiado.")


def detener_busqueda():
    global cronometro_activo
    if cronometro_activo:
        cronometro_activo = False
        estado_label.configure(text="⏹ Búsqueda detenida por el usuario.")
        log("🛑 La búsqueda ha sido detenida manualmente.")
    else:
        log("⚠️ No hay una búsqueda en curso.")


def on_closing():
    if messagebox.askokcancel("Salir", "¿Seguro que quieres cerrar la aplicación?"):
        app.destroy()


app.protocol("WM_DELETE_WINDOW", on_closing)

# === Botones principales ===
frame_botones = ctk.CTkFrame(app, corner_radius=20)
frame_botones.pack(pady=10)

ctk.CTkButton(frame_botones, text="1. Seleccionar Audio", command=seleccionar_audio,
              width=100, corner_radius=12).pack(side="left", padx=10)
ctk.CTkButton(frame_botones, text="2. Ingresar Frases", command=pedir_frases_clave,
              width=100, corner_radius=12).pack(side="left", padx=10)
ctk.CTkButton(frame_botones, text="3. Buscar Coincidencias",
              command=lambda: threading.Thread(
                  target=detectar_fragmentos, daemon=True).start(),
              width=100, corner_radius=12).pack(side="left", padx=10)
ctk.CTkButton(frame_botones, text="4. Exportar a Word", command=exportar_a_word,
              width=100, corner_radius=12).pack(side="left", padx=10)
ctk.CTkButton(frame_botones, text="5. Limpiar", command=limpiar_todo,
              width=100, corner_radius=12).pack(side="left", padx=10)
ctk.CTkButton(frame_botones, text="6. Detener búsqueda", command=detener_busqueda,
              width=100, corner_radius=12).pack(side="left", padx=10)


# === Ejecutar la app ===
app.mainloop()
# pyinstaller --noconfirm --onefile --windowed --icon=icons/ico.ico transcribir_16.py

'''pyinstaller --noconfirm --onefile --windowed --icon=icons/ico.ico \
--add-data "icons;icons" \
--add-data "whisper;whisper" \
--add-data "pyannote_env;pyannote_env" \
analizador_audio.py
'''
